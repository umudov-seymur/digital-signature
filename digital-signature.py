import hashlib
from docx import Document
from datetime import datetime

def calculate_sha256(data):
    sha256_hash = hashlib.sha256()
    sha256_hash.update(data)
    return sha256_hash.hexdigest()

def sign_document(file_path):
    # Sənədi aç
    doc = Document(file_path)

    # Sənədin içindəkiləri oxu
    document_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    # Sənəd içindəkiləri sha256 ilə şifrələ
    document_hash = calculate_sha256(document_content.encode())

    # Imza yarat
    signature = f"Digital Signature: {document_hash}\nTimestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

    # Imzanı sənədə əlavə et
    doc.add_paragraph(signature)

    # İmzalı sənədi yadda saxla
    signed_file_path = "signed_document.docx"
    doc.save(signed_file_path)

    print(f"Sənəd digital olaraq imzalandı və '{signed_file_path}' olaraq yadda saxlanıldı.")

def verify_signature(file_path):
    # Sənədi aç
    doc = Document(file_path)

    # İmzanın var olub olmadığını yoxla
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Digital Signature"):
            signature_line = paragraph.text
            signature, timestamp = map(str.strip, signature_line.split('\n'))

            # Sənəd içindəkiləri al
            document_content = "\n".join([para.text for para in doc.paragraphs if para.text != signature_line])

            # Sənəd içindəkiləri SHA-256 ilə hash'le
            document_hash = calculate_sha256(document_content.encode())

            # Imzayı və tarixi kontrol elə
            if signature == f"Digital Signature: {document_hash}":
                print(f"Digital imza doğrulandı. Sənəd güvənli. Imza tarixi: {timestamp.split(': ')[1]}")
            else:
                print("Digital imza doğrulanmadı. Sənəd güvənli deyil.")
            return

    print("Digital imza tapılmadı. Sənəd güvənli deyil.")

# Sənədi imzala
sign_document("test.docx")

# İmzalanmış bir sənədin rəqəmsal imzasını və tarixini yoxla
verify_signature("signed_document.docx")
